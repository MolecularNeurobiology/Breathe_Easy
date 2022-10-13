#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Mon Oct  3 14:23:29 2022

@author: S_Barnett
"""

import os
import pandas as pd
from tkinter import filedialog as fd
from docx import Document
from docx.shared import Inches

input_folder=fd.askdirectory()

#%%
#Generates a list of all files within the directory folder selected above as input_folder.
fileList=os.walk(input_folder)

#Stores filepaths for each of the 3 configuration files used by STAGG.
for d,p,f in fileList:
    for fileName in f:
        if "graph_config" in fileName:
            graphPath=os.path.join(d,fileName)
        if "variable_config" in fileName:
            variablePath=os.path.join(d,fileName)
        if "other_config" in fileName:
            optionalPath=os.path.join(d,fileName)

#Reads in the tables from the configuration files that were found above.
graphTable=pd.read_csv(graphPath)
variableTable=pd.read_csv(variablePath)
optionalTable=pd.read_csv(optionalPath)

#Variable config requires filtering to remove variables that are not part of the primary model. The below code
# removes all unused variables from the table.
relevantVariables=variableTable[(variableTable["Independent"] == 1) | (variableTable["Dependent"] == 1) | (variableTable["Covariate"] == 1)]

#%%
README=Document()

sections = README.sections
for section in sections:
    section.top_margin=Inches(0.5)
    section.bottom_margin=Inches(0.5)
    section.left_margin=Inches(0.5)
    section.right_margin=Inches(0.5)

README.add_heading('STAGG Settings Summary', 0)

README.add_heading('Notes', 0)
README.add_paragraph('Add notes on the output of this run here', style='List Number')

README.add_heading('Graphing settings', 0)
GT=README.add_table(graphTable.shape[0]+1, graphTable.shape[1])

for j in range(graphTable.shape[-1]):
    GT.cell(0,j).text=graphTable.columns[j]
    
for i in range(graphTable.shape[0]):
    for j in range(graphTable.shape[-1]):
        GT.cell(i+1,j).text = str(graphTable.values[i,j])
        
GT.style="Light Grid"

README.add_heading('Variable settings', 0)
RV=README.add_table(relevantVariables.shape[0]+1, relevantVariables.shape[1])

for j in range(relevantVariables.shape[-1]):
    RV.cell(0,j).text=relevantVariables.columns[j]
    
for i in range(relevantVariables.shape[0]):
    for j in range(relevantVariables.shape[-1]):
        RV.cell(i+1,j).text = str(relevantVariables.values[i,j])

RV.style="Light Grid"
RV.allow_autofit=False
RV.autofit = False

col = RV.columns[0]
for cell in col.cells:
    cell.width=Inches(0.5)

README.add_heading('Optional settings', 0)
OT=README.add_table(optionalTable.shape[0]+1, optionalTable.shape[1])

for j in range(optionalTable.shape[-1]):
    OT.cell(0,j).text=optionalTable.columns[j]
    
for i in range(optionalTable.shape[0]):
    for j in range(optionalTable.shape[-1]):
        OT.cell(i+1,j).text = str(optionalTable.values[i,j])
        
OT.style="Light Grid"

README.save('1_README.docx')