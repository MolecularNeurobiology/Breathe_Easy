#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Mon Oct  3 14:23:29 2022

@author: S_Barnett
"""

from glob import glob
import os
import pandas as pd
from tkinter import filedialog as fd
from docx import Document
from docx.shared import Inches

def retrieve_settings(input_folder: str, settings_key: str) -> pd.DataFrame:
    """Search for and read a given settings file"""
    # Search for path matching string regex
    settings_paths = glob(os.path.join(input_folder, f"*{settings_key}_config*"))
    if len(settings_paths) > 1:
        raise RuntimeError(f"Too many {settings_key} settings in the input folder {input_folder}")
    elif len(settings_paths) == 0:
        raise FileNotFoundError(f"No {settings_key} settings found in the input folder {input_folder}")
    # Read the file and return the dataframe
    return pd.read_csv(settings_paths[0])

def prepare_document(document: Document):
    """Apply general formatting to document"""
    for section in document.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    document.add_heading('STAGG Settings Summary', 0)

def add_notes(document: Document):
    """Add notes section to document"""
    document.add_heading('Notes', 0)
    document.add_paragraph('Add notes on the output of this run here', style='List Number')

def add_table(document: Document, table_key: str, dataframe: pd.DataFrame):
    """Add dataframe contents to document"""
    document.add_heading(f'{table_key.capitalize()} Settings', 0)
    new_table = document.add_table(dataframe.shape[0]+1, dataframe.shape[1])

    for j in range(dataframe.shape[-1]):
        new_table.cell(0, j).text = dataframe.columns[j]
        
    for i in range(dataframe.shape[0]):
        for j in range(dataframe.shape[-1]):
            new_table.cell(i+1, j).text = str(dataframe.values[i, j])

    new_table.style = "Light Grid"
    return new_table

def add_graph_table(document: Document, graph_data: pd.DataFrame):
    """Add graph dataframe contents to document"""
    add_table(document, table_key='graph', dataframe=graph_data)

def add_optional_table(document: Document, optional_data: pd.DataFrame):
    """Add optional dataframe contents to document"""
    add_table(document, table_key='Optional', dataframe=optional_data)

def add_variable_table(document: Document, variable_data: pd.DataFrame):
    """Add variable dataframe contents to document"""

    # Variable config requires filtering to remove variables that are not part of the primary model.
    # The below code removes all unused variables from the table.
    relevant_variables = variable_data[(variable_data["Independent"] == 1) |
                                       (variable_data["Dependent"] == 1) |
                                       (variable_data["Covariate"] == 1)]

    RV = add_table(document, table_key='variable', dataframe=relevant_variables)

    # Some extra formatting
    RV.allow_autofit = False
    RV.autofit = False
    col = RV.columns[0]
    for cell in col.cells:
        cell.width = Inches(0.5)

def get_data_and_generate_document(input_folder: str, output_dir: str = ""):
    # Read in the tables for each of the 3 configuration files used by STAGG
    graph_data = retrieve_settings(input_folder, 'graph')
    variable_data = retrieve_settings(input_folder, 'variable')
    optional_data = retrieve_settings(input_folder, 'other')
    generate_document(variable_data, graph_data, optional_data, output_dir)

def generate_document(variable_data: pd.DataFrame, graph_data: pd.DataFrame,
                      optional_data: pd.DataFrame, output_dir: str = ""):
    # Create document
    README = Document()

    # Format document and add title header
    prepare_document(README)

    # Add all sections
    add_notes(README)
    add_graph_table(README, graph_data)
    add_variable_table(README, variable_data)
    add_optional_table(README, optional_data)

    README.save(os.path.join(output_dir, 'README.docx'))

if __name__ == '__main__':
    input_folder = fd.askdirectory()  # Get input from user
    get_data_and_generate_document(input_folder)

